from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bidi.algorithm import get_display
import arabic_reshaper

def ar(text):

    reshaped_text = arabic_reshaper.reshape(
        str(text)
    )

    return get_display(
        reshaped_text
    )

def generate_word_report(
    file_name,
    report_text
):

    doc = Document()


    styles = doc.styles

    styles["Normal"].paragraph_format.alignment = (
        WD_PARAGRAPH_ALIGNMENT.RIGHT
    )



    title = doc.add_heading(
        "تقرير مشروع",
        level=1
    )

    title.alignment = (
        WD_PARAGRAPH_ALIGNMENT.RIGHT
    )


    title_pPr = title._element.get_or_add_pPr()

    title_bidi = OxmlElement("w:bidi")

    title_bidi.set(
        qn("w:val"),
        "1"
    )

    title_pPr.append(
        title_bidi
    )



    title.alignment = (
        WD_PARAGRAPH_ALIGNMENT.RIGHT
    )


    for line in report_text.split("\n"):

        p = doc.add_paragraph()

        p.alignment = (
            WD_PARAGRAPH_ALIGNMENT.RIGHT
        )

        pPr = p._element.get_or_add_pPr()

        bidi = OxmlElement("w:bidi")

        bidi.set(
            qn("w:val"),
            "1"
        )

        pPr.append(bidi)

        run = p.add_run(
            line.strip()
        )


        run.font.size = Pt(12)


    section = doc.sections[0]

    section.right_margin = 500000
    section.left_margin = 500000

    doc.save(
        file_name
    )